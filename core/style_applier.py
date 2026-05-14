"""
样式应用器
将样式模板应用到文档
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from typing import Dict
from config.style_templates import StyleTemplate


class StyleApplier:
    """样式应用器"""
    
    def __init__(self, document: Document, style_template: StyleTemplate):
        """
        初始化样式应用器
        
        Args:
            document: DOCX文档对象
            style_template: 样式模板
        """
        self.document = document
        self.style_template = style_template
        
    def apply_all_styles(self):
        """应用所有样式"""
        # 设置页面边距
        self._set_page_margins()
        
        # 应用标题样式
        self._apply_title_style()
        
        # 应用章节标题样式
        self._apply_heading_styles()
        
        # 应用正文样式
        self._apply_normal_style()
        
        # 应用表格样式
        self._apply_table_styles()
        
    def _set_page_margins(self):
        """设置页面边距"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)
    
    def _apply_title_style(self):
        """应用标题样式"""
        title_style = self.style_template.get_title_style()
        
        # 查找文档第一个段落作为标题（如果存在）
        if self.document.paragraphs:
            first_para = self.document.paragraphs[0]
            # 简单判断：如果第一段较短且居中，可能是标题
            if len(first_para.text) < 100:
                self._apply_paragraph_style(first_para, title_style)
    
    def _apply_heading_styles(self):
        """应用章节标题样式"""
        heading_styles = self.style_template.get_heading_styles()
        
        for paragraph in self.document.paragraphs:
            if not paragraph.text.strip():
                continue
                
            # 检测是否为标题
            heading_level = self._detect_heading_level(paragraph)
            
            if heading_level > 0:
                style_key = f'heading{heading_level}'
                if style_key in heading_styles:
                    self._apply_paragraph_style(paragraph, heading_styles[style_key])
    
    def _apply_normal_style(self):
        """应用正文样式"""
        normal_style = self.style_template.get_normal_style()
        
        for paragraph in self.document.paragraphs:
            if not paragraph.text.strip():
                continue
            
            # 跳过已应用标题样式的段落
            if self._detect_heading_level(paragraph) > 0:
                continue
            
            # 跳过可能的标题（第一段）
            if paragraph == self.document.paragraphs[0] and len(paragraph.text) < 100:
                continue
            
            self._apply_paragraph_style(paragraph, normal_style)
    
    def _apply_table_styles(self):
        """应用表格样式"""
        table_style = self.style_template.get_table_style()
        
        for table in self.document.tables:
            self._format_table(table, table_style)
    
    def _apply_paragraph_style(self, paragraph, style_config: Dict):
        """
        应用段落样式
        
        Args:
            paragraph: 段落对象
            style_config: 样式配置
        """
        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        
        if 'alignment' in style_config:
            paragraph_format.alignment = style_config['alignment']
        
        if 'space_before' in style_config:
            paragraph_format.space_before = style_config['space_before']
        
        if 'space_after' in style_config:
            paragraph_format.space_after = style_config['space_after']
        
        if 'first_line_indent' in style_config:
            paragraph_format.first_line_indent = style_config['first_line_indent']
        
        if 'line_spacing' in style_config:
            paragraph_format.line_spacing = style_config['line_spacing']
        
        # 设置字体格式
        for run in paragraph.runs:
            if 'font_name' in style_config:
                run.font.name = style_config['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), style_config['font_name'])
            
            if 'font_size' in style_config:
                run.font.size = style_config['font_size']
            
            if 'bold' in style_config:
                run.bold = style_config['bold']
            
            if 'color' in style_config:
                run.font.color.rgb = style_config['color']
    
    def _format_table(self, table, style_config: Dict):
        """
        格式化表格
        
        Args:
            table: 表格对象
            style_config: 表格样式配置
        """
        # 设置表格边框
        self._set_table_borders(table, style_config.get('border_width', Pt(1)))
        
        # 格式化表头
        if table.rows:
            header_row = table.rows[0]
            self._format_table_row(header_row, style_config, is_header=True)
            
            # 格式化其他行
            for row in table.rows[1:]:
                self._format_table_row(row, style_config, is_header=False)
    
    def _format_table_row(self, row, style_config: Dict, is_header: bool = False):
        """格式化表格行"""
        for cell in row.cells:
            # 设置单元格内段落样式
            for paragraph in cell.paragraphs:
                paragraph_format = paragraph.paragraph_format
                
                if 'alignment' in style_config:
                    paragraph_format.alignment = style_config['alignment']
                
                for run in paragraph.runs:
                    if 'font_size' in style_config:
                        run.font.size = style_config['font_size']
                    
                    if is_header:
                        run.bold = True
    
    def _set_table_borders(self, table, border_width: Pt):
        """设置表格边框"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._element
        
        # 创建边框元素
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(int(border_width.pt * 8)))  # 转换为eighth-points
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblPr.append(tblBorders)
    
    def _detect_heading_level(self, paragraph) -> int:
        """
        检测段落的标题级别
        
        Args:
            paragraph: 段落对象
            
        Returns:
            标题级别（0表示不是标题）
        """
        # 方法1：通过样式名称判断
        if paragraph.style:
            style_name = paragraph.style.name.lower()
            import re
            match = re.search(r'heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))
            
            if 'title' in style_name:
                return 1
        
        # 方法2：通过格式特征判断
        text = paragraph.text.strip()
        if not text:
            return 0
        
        # 检查是否加粗
        is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
        
        # 检查文本长度（标题通常较短）
        if len(text) > 100:
            return 0
        
        # 检查是否包含序号（如：一、二、三 或 1. 2. 3.）
        import re
        has_number = bool(re.match(r'^[一二三四五六七八九十\d]+[、.]', text))
        
        if is_bold and (has_number or len(text) < 50):
            return 2  # 默认为二级标题
        
        return 0
