"""
文档分析器
提取文档的结构信息和文本内容
"""
from docx import Document
from docx.shared import Pt, RGBColor
from typing import Dict, List, Tuple
import re


class DocumentAnalyzer:
    """文档分析器"""
    
    def __init__(self, file_path: str):
        """
        初始化文档分析器
        
        Args:
            file_path: DOCX文件路径
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self.text_content = ""
        self.structure_info = {}
        self.paragraphs_info = []
        
    def analyze(self) -> Dict:
        """
        分析文档，提取内容和结构信息
        
        Returns:
            包含文本内容和结构信息的字典
        """
        # 提取文本内容
        self._extract_text()
        
        # 分析段落结构
        self._analyze_paragraphs()
        
        # 检测结构特征
        self._detect_structure_features()
        
        return {
            'text_content': self.text_content,
            'structure_info': self.structure_info,
            'paragraphs_info': self.paragraphs_info,
            'total_paragraphs': len(self.document.paragraphs),
            'total_tables': len(self.document.tables),
            'heading_count': self._count_headings()
        }
    
    def _extract_text(self):
        """提取文档所有文本内容"""
        texts = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text.strip())
        
        self.text_content = '\n'.join(texts)
    
    def _analyze_paragraphs(self):
        """分析每个段落的样式和结构"""
        paragraphs_info = []
        
        for idx, paragraph in enumerate(self.document.paragraphs):
            info = {
                'index': idx,
                'text': paragraph.text,
                'style': paragraph.style.name if paragraph.style else 'Normal',
                'is_heading': self._is_heading(paragraph),
                'heading_level': self._get_heading_level(paragraph),
                'font_size': self._get_font_size(paragraph),
                'is_bold': self._is_bold(paragraph),
                'alignment': paragraph.alignment,
                'has_table': False
            }
            paragraphs_info.append(info)
        
        self.paragraphs_info = paragraphs_info
    
    def _detect_structure_features(self):
        """检测文档的结构特征"""
        features = {
            'has_signature_section': self._has_signature_section(),
            'has_party_info': self._has_party_info(),
            'has_clauses': self._has_clauses(),
            'formal_language': self._has_formal_language(),
            'has_steps': self._has_steps(),
            'has_warnings': self._has_warnings(),
            'has_specifications': self._has_specifications(),
            'has_toc': self._has_toc(),
            'has_abstract': self._has_abstract(),
            'has_sections': self._has_sections(),
            'has_references': self._has_references(),
            'has_charts': len(self.document.tables) > 0 or self._has_images(),
            'has_background': self._has_background(),
            'has_plan': self._has_plan(),
            'has_budget': self._has_budget(),
            'has_timeline': self._has_timeline()
        }
        
        self.structure_info = features
    
    def _is_heading(self, paragraph) -> bool:
        """判断是否为标题"""
        style_name = paragraph.style.name.lower() if paragraph.style else ''
        return 'heading' in style_name or 'title' in style_name
    
    def _get_heading_level(self, paragraph) -> int:
        """获取标题级别"""
        if not paragraph.style:
            return 0
        
        style_name = paragraph.style.name.lower()
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        
        if 'title' in style_name:
            return 1
        
        return 0
    
    def _get_font_size(self, paragraph) -> float:
        """获取段落字体大小"""
        if paragraph.runs:
            for run in paragraph.runs:
                if run.font.size:
                    return run.font.size.pt
        return 12.0  # 默认字体大小
    
    def _is_bold(self, paragraph) -> bool:
        """判断段落是否加粗"""
        if paragraph.runs:
            for run in paragraph.runs:
                if run.bold:
                    return True
        return False
    
    def _count_headings(self) -> int:
        """统计标题数量"""
        count = 0
        for para in self.document.paragraphs:
            if self._is_heading(para):
                count += 1
        return count
    
    def _has_signature_section(self) -> bool:
        """检测是否有签署区域"""
        signature_keywords = ['签字', '盖章', '签署', 'signature', 'signed']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in signature_keywords)
    
    def _has_party_info(self) -> bool:
        """检测是否有甲乙双方信息"""
        party_keywords = ['甲方', '乙方', 'party a', 'party b', '买方', '卖方']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in party_keywords)
    
    def _has_clauses(self) -> bool:
        """检测是否有条款编号"""
        clause_pattern = r'第[一二三四五六七八九十\d]+条'
        return bool(re.search(clause_pattern, self.text_content))
    
    def _has_formal_language(self) -> bool:
        """检测是否使用正式语言"""
        formal_words = ['兹', '鉴于', '特此', '依照', '根据', '规定']
        text_lower = self.text_content.lower()
        count = sum(1 for word in formal_words if word in text_lower)
        return count >= 2
    
    def _has_steps(self) -> bool:
        """检测是否有步骤说明"""
        step_patterns = [r'步骤\s*\d+', r'第一步', r'第二步', r'\d+\.\s*']
        return any(re.search(pattern, self.text_content) for pattern in step_patterns)
    
    def _has_warnings(self) -> bool:
        """检测是否有警告提示"""
        warning_keywords = ['注意', '警告', ' caution', 'warning', '重要']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in warning_keywords)
    
    def _has_specifications(self) -> bool:
        """检测是否有技术规格"""
        spec_keywords = ['规格', '参数', 'technical', 'specification', '型号']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in spec_keywords)
    
    def _has_toc(self) -> bool:
        """检测是否有目录"""
        toc_keywords = ['目录', 'contents', 'table of contents']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in toc_keywords)
    
    def _has_abstract(self) -> bool:
        """检测是否有摘要"""
        abstract_keywords = ['摘要', 'abstract', '概要', '概述']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in abstract_keywords)
    
    def _has_sections(self) -> bool:
        """检测是否有章节"""
        section_keywords = ['第一章', '第二章', 'section', 'chapter']
        return any(keyword in self.text_content for keyword in section_keywords)
    
    def _has_references(self) -> bool:
        """检测是否有参考文献"""
        ref_keywords = ['参考文献', 'references', '引用', 'bibliography']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in ref_keywords)
    
    def _has_images(self) -> bool:
        """检测是否有图片"""
        # python-docx中图片检测较为复杂，这里简化处理
        return False
    
    def _has_background(self) -> bool:
        """检测是否有背景介绍"""
        bg_keywords = ['背景', 'background', '项目背景', '研究背景']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in bg_keywords)
    
    def _has_plan(self) -> bool:
        """检测是否有方案内容"""
        plan_keywords = ['方案', 'plan', '实施', 'solution']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in plan_keywords)
    
    def _has_budget(self) -> bool:
        """检测是否有预算"""
        budget_keywords = ['预算', 'budget', '费用', '成本']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in budget_keywords)
    
    def _has_timeline(self) -> bool:
        """检测是否有时间计划"""
        timeline_keywords = ['进度', 'timeline', '计划', 'schedule', '时间安排']
        text_lower = self.text_content.lower()
        return any(keyword in text_lower for keyword in timeline_keywords)
