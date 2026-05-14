"""
测试文档生成器
用于创建测试用的DOCX文档
"""
from docx import Document
from docx.shared import Pt, Inches
import os


def create_sample_contract(file_path: str):
    """创建示例合同文档"""
    doc = Document()
    
    # 标题
    title = doc.add_paragraph('产品销售合同')
    title.style = 'Title'
    
    # 合同编号
    doc.add_paragraph('合同编号：HT-2024-001')
    
    # 甲方乙方
    doc.add_paragraph('甲方（买方）：某某科技有限公司')
    doc.add_paragraph('乙方（卖方）：某某贸易有限公司')
    
    # 第一条
    h1 = doc.add_paragraph('第一条 产品规格')
    h1.style = 'Heading 1'
    
    doc.add_paragraph('乙方向甲方提供以下产品：')
    doc.add_paragraph('1. 产品名称：XXX设备')
    doc.add_paragraph('2. 数量：100台')
    doc.add_paragraph('3. 单价：人民币5000元')
    
    # 第二条
    h2 = doc.add_paragraph('第二条 付款方式')
    h2.style = 'Heading 1'
    
    doc.add_paragraph('甲方应按以下方式付款：')
    doc.add_paragraph('1. 合同签订后支付30%预付款')
    doc.add_paragraph('2. 交货前支付60%货款')
    doc.add_paragraph('3. 验收合格后支付10%尾款')
    
    # 第三条
    h3 = doc.add_paragraph('第三条 交货日期')
    h3.style = 'Heading 1'
    
    doc.add_paragraph('乙方应于合同签订后30日内交货。')
    
    # 第四条
    h4 = doc.add_paragraph('第四条 违约责任')
    h4.style = 'Heading 1'
    
    doc.add_paragraph('任何一方违反本合同约定，应承担违约责任。')
    
    # 签署区域
    doc.add_paragraph('\n\n甲方（盖章）：_______________    乙方（盖章）：_______________')
    doc.add_paragraph('签字：_______________              签字：_______________')
    doc.add_paragraph('日期：_______________              日期：_______________')
    
    doc.save(file_path)
    print(f"已创建示例合同: {file_path}")


def create_sample_manual(file_path: str):
    """创建示例说明书文档"""
    doc = Document()
    
    # 标题
    title = doc.add_paragraph('XXX设备使用说明书')
    title.style = 'Title'
    
    # 目录
    doc.add_paragraph('目录')
    doc.add_paragraph('一、产品简介')
    doc.add_paragraph('二、技术参数')
    doc.add_paragraph('三、安装步骤')
    doc.add_paragraph('四、使用说明')
    doc.add_paragraph('五、注意事项')
    doc.add_paragraph('六、故障排除')
    
    # 第一章
    h1 = doc.add_paragraph('一、产品简介')
    h1.style = 'Heading 1'
    
    doc.add_paragraph('本产品是一款高性能的XXX设备，具有以下特点：')
    doc.add_paragraph('1. 高效节能')
    doc.add_paragraph('2. 操作简便')
    doc.add_paragraph('3. 安全可靠')
    
    # 第二章
    h2 = doc.add_paragraph('二、技术参数')
    h2.style = 'Heading 1'
    
    doc.add_paragraph('主要技术指标如下：')
    doc.add_paragraph('• 功率：1000W')
    doc.add_paragraph('• 电压：220V')
    doc.add_paragraph('• 尺寸：500mm x 300mm x 200mm')
    doc.add_paragraph('• 重量：15kg')
    
    # 第三章
    h3 = doc.add_paragraph('三、安装步骤')
    h3.style = 'Heading 1'
    
    doc.add_paragraph('请按照以下步骤进行安装：')
    doc.add_paragraph('步骤1：开箱检查配件是否齐全')
    doc.add_paragraph('步骤2：选择平整的安装位置')
    doc.add_paragraph('步骤3：连接电源线')
    doc.add_paragraph('步骤4：开机测试')
    
    # 第四章
    h4 = doc.add_paragraph('四、使用说明')
    h4.style = 'Heading 1'
    
    doc.add_paragraph('1. 按下电源开关启动设备')
    doc.add_paragraph('2. 设置工作参数')
    doc.add_paragraph('3. 开始运行')
    
    # 第五章
    h5 = doc.add_paragraph('五、注意事项')
    h5.style = 'Heading 1'
    
    doc.add_paragraph('警告：使用前请仔细阅读本说明书！')
    doc.add_paragraph('注意：请勿在潮湿环境中使用')
    doc.add_paragraph('重要：定期维护保养可延长使用寿命')
    
    # 第六章
    h6 = doc.add_paragraph('六、故障排除')
    h6.style = 'Heading 1'
    
    doc.add_paragraph('如遇到问题，请参考以下解决方案：')
    doc.add_paragraph('问题1：无法启动 - 检查电源连接')
    doc.add_paragraph('问题2：异常噪音 - 检查零部件是否松动')
    
    doc.save(file_path)
    print(f"已创建示例说明书: {file_path}")


def create_sample_report(file_path: str):
    """创建示例报告文档"""
    doc = Document()
    
    # 标题
    title = doc.add_paragraph('市场调研分析报告')
    title.style = 'Title'
    
    # 摘要
    h_abstract = doc.add_paragraph('摘要')
    h_abstract.style = 'Heading 1'
    
    doc.add_paragraph('本报告对当前市场状况进行了全面分析，旨在为决策提供参考依据。')
    
    # 引言
    h_intro = doc.add_paragraph('一、引言')
    h_intro.style = 'Heading 1'
    
    doc.add_paragraph('随着经济的发展，市场需求不断变化。本研究旨在了解当前市场趋势。')
    
    # 研究方法
    h_method = doc.add_paragraph('二、研究方法')
    h_method.style = 'Heading 1'
    
    doc.add_paragraph('本次调研采用问卷调查和深度访谈相结合的方式。')
    
    # 调查结果
    h_result = doc.add_paragraph('三、调查结果')
    h_result.style = 'Heading 1'
    
    doc.add_paragraph('调查数据显示：')
    doc.add_paragraph('1. 市场规模持续增长')
    doc.add_paragraph('2. 消费者偏好发生变化')
    doc.add_paragraph('3. 竞争格局趋于激烈')
    
    # 数据分析
    h_analysis = doc.add_paragraph('四、数据分析')
    h_analysis.style = 'Heading 1'
    
    doc.add_paragraph('通过统计分析，我们发现以下规律...')
    
    # 结论
    h_conclusion = doc.add_paragraph('五、结论')
    h_conclusion.style = 'Heading 1'
    
    doc.add_paragraph('基于以上分析，得出以下结论...')
    
    # 建议
    h_suggestion = doc.add_paragraph('六、建议')
    h_suggestion.style = 'Heading 1'
    
    doc.add_paragraph('针对研究发现，提出以下建议：')
    doc.add_paragraph('1. 加强产品创新')
    doc.add_paragraph('2. 优化营销策略')
    doc.add_paragraph('3. 提升服务质量')
    
    # 参考文献
    h_ref = doc.add_paragraph('参考文献')
    h_ref.style = 'Heading 1'
    
    doc.add_paragraph('[1] 张三. 市场分析方法[M]. 北京: 出版社, 2023.')
    doc.add_paragraph('[2] 李四. 消费者行为研究[J]. 经济学报, 2024(1): 10-20.')
    
    doc.save(file_path)
    print(f"已创建示例报告: {file_path}")


def create_all_samples(output_dir: str = './test_docs'):
    """创建所有示例文档"""
    os.makedirs(output_dir, exist_ok=True)
    
    create_sample_contract(os.path.join(output_dir, '示例合同.docx'))
    create_sample_manual(os.path.join(output_dir, '示例说明书.docx'))
    create_sample_report(os.path.join(output_dir, '示例报告.docx'))
    
    print(f"\n所有示例文档已创建到: {output_dir}")


if __name__ == '__main__':
    create_all_samples()
